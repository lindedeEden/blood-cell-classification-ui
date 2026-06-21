# -*- coding: utf-8 -*-
"""Generate Word and Excel facilitator observation forms."""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

THIN = Side(style='thin', color='999999')
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
HEADER_FILL = PatternFill('solid', fgColor='4472C4')
HEADER_FONT = Font(bold=True, color='FFFFFF', size=11)
TITLE_FONT = Font(bold=True, size=14)
SECTION_FONT = Font(bold=True, size=12)
NORMAL = Font(size=11)


def style_header_row(ws, row, cols):
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = BORDER


def apply_table_borders(ws, min_row, max_row, min_col, max_col):
    for r in range(min_row, max_row + 1):
        for c in range(min_col, max_col + 1):
            ws.cell(row=r, column=c).border = BORDER
            ws.cell(row=r, column=c).alignment = Alignment(vertical='center', wrap_text=True)


def set_col_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


def kv_table(ws, start_row, items, col1_w=22, col2_w=50):
    r = start_row
    for label, _ in items:
        ws.cell(row=r, column=1, value=label).font = Font(bold=True, size=11)
        ws.cell(row=r, column=2, value='')
        r += 1
    apply_table_borders(ws, start_row, r - 1, 1, 2)
    set_col_widths(ws, [col1_w, col2_w])
    return r


def build_excel(path):
    wb = Workbook()

    # --- Sheet: 填寫說明 ---
    ws0 = wb.active
    ws0.title = '填寫說明'
    ws0['A1'] = '林口長庚血球型態分類軟體 — 改善後成效調查'
    ws0['A1'].font = TITLE_FONT
    ws0['A2'] = '主持人觀察紀錄表（受試者不可見）｜全程約 40 分鐘'
    ws0['A4'] = '一、研究目的'
    ws0['A4'].font = SECTION_FONT
    ws0['A5'] = '驗證改善後模擬介面，是否在五項高提及率痛點上，比現行軟體更能支援醫檢師完成血球型態閱片相關判斷與作業。'
    ws0['A7'] = '二、整體流程與時間'
    ws0['A7'].font = SECTION_FONT
    flow_lines = [
        ['階段', '內容', '時間'],
        ['階段一', '簡介與受試者同意', '3 min'],
        ['導覽', '原型介面導覽介紹', '5 min'],
        ['階段二', '五項情境操作（成功率/耗時/SEQ）', '20 min'],
        ['階段三', 'SUS 改善前後比較', '10 min'],
    ]
    for i, row in enumerate(flow_lines, 8):
        for c, v in enumerate(row, 1):
            ws0.cell(row=i, column=c, value=v)
    ws0['A14'] = 'SEQ：每情境結束後立即施測，7點量表（1=非常困難～7=非常容易）'
    ws0['A15'] = 'SUS：先填改良版，再回想現行版；10題×5點，後算總分(0-100)'
    set_col_widths(ws0, [80])

    # --- Sheet: 階段一同意 ---
    ws_consent = wb.create_sheet('階段一同意')
    ws_consent['A1'] = '階段一：簡介與受試者同意'
    ws_consent['A1'].font = SECTION_FONT
    ws_consent['A2'] = '簡介：聆聽簡介→介面導覽→5項情境→SUS問卷，約40分鐘；假資料模擬，不影響真實LIS。'
    ws_consent['A3'] = '報告以P1～P5編號呈現；參與自願可隨時退出。'
    kv_table(ws_consent, 5, [
        ('測試日期', ''), ('受試者編號', ''), ('職級', '□醫檢師 □住院醫師 □其他'),
        ('年資', '□<1年 □1-3年 □3-5年 □>5年'), ('主持人', ''), ('觀察員', ''),
        ('受試者簽名', ''), ('說明人簽名', ''), ('同意時間', ''),
    ])

    # --- Sheet: 基本資料 ---
    ws1 = wb.create_sheet('基本資料')
    ws1['A1'] = '階段二：基本紀錄與流程'
    ws1['A1'].font = SECTION_FONT
    items = [
        ('測試日期', ''), ('場次編號', ''), ('受試者編號', ''),
        ('職級', '□ 醫檢師  □ 住院醫師  □ 其他'),
        ('年資', '□ <1年  □ 1-3年  □ 3-5年  □ >5年'),
        ('主持人', ''), ('觀察員', ''),
        ('測試開始時間', ''), ('測試結束時間', ''),
        ('備註', ''),
    ]
    kv_table(ws1, 3, items)

    ws1['A15'] = '五項情境總覽（成功率/耗時/SEQ）'
    ws1['A15'].font = SECTION_FONT
    headers = ['序', '情境', '建議模式', '成功定義', '耗時', '成功', 'SEQ']
    for c, h in enumerate(headers, 1):
        ws1.cell(row=16, column=c, value=h)
    style_header_row(ws1, 16, len(headers))
    flow = [
        ['1', '推片確認資訊缺失', '受試者自選', '3筆FU全對未選6301', '', '□是□否', ''],
        ['2', '位置／機台缺失', '實體作業', '3歸位+2機台', '', '□是□否', ''],
        ['3', '歷史報告缺失', '數位閱片', '3筆前次+關聯', '', '□是□否', ''],
        ['4~5', '異常／留單警示', '數位閱片', '異常5/留單3', '', '□是□否', ''],
    ]
    for i, row in enumerate(flow, 17):
        for c, v in enumerate(row, 1):
            ws1.cell(row=i, column=c, value=v)
    apply_table_borders(ws1, 16, 20, 1, 7)
    set_col_widths(ws1, [6, 22, 12, 22, 10, 10, 8])

    # --- Sheet: 情境一 ---
    ws2 = wb.create_sheet('情境一')
    ws2['A1'] = '情境一：需推片確認資訊缺失'
    ws2['A1'].font = SECTION_FONT
    ws2['A2'] = '任務：找出所有仍需「需拉片確認(Follow-up)」的檢體，口頭說明編號與判斷依據'
    kv_table(ws2, 4, [
        ('開始時間', ''), ('完成時間', ''), ('耗時(秒)', ''),
        ('是否說明判斷依據', '□是 □否'),
        ('是否誤選 H5080706301', '□是 □否'),
        ('正確數', '/ 3'),
    ])
    ws2['A12'] = '標準答案對照'
    ws2['A12'].font = Font(bold=True, size=11)
    h1 = ['檢體編號', '是否應選', '判斷依據', '受試者是否提及', '正確']
    for c, h in enumerate(h1, 1):
        ws2.cell(row=13, column=c, value=h)
    style_header_row(ws2, 13, len(h1))
    s1 = [
        ['H5080706280', '應選', 'Follow-up待辦；歸位090-2', '', ''],
        ['H5080720696', '應選', 'Follow-up待辦；歸位125-12', '', ''],
        ['H5080721201', '應選', 'AI+Follow-up雙旗標；歸位208-6', '', ''],
        ['H5080706301', '不可選', '陷阱：Follow-up已完成', '', ''],
        ['其餘6筆', '不可選', '非待辦Follow-up', '', ''],
    ]
    for i, row in enumerate(s1, 14):
        for c, v in enumerate(row, 1):
            ws2.cell(row=i, column=c, value=v)
    apply_table_borders(ws2, 13, 18, 1, 5)
    ws2['A20'] = 'SEQ — 情境一'
    ws2['A20'].font = SECTION_FONT
    kv_table(ws2, 21, [
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'),
        ('受試者理由', ''),
    ])
    set_col_widths(ws2, [16, 10, 32, 18, 8])

    # --- Sheet: 情境二 ---
    ws3 = wb.create_sheet('情境二')
    ws3['A1'] = '情境二：檢體實體位置／機台位置缺失'
    ws3['A1'].font = SECTION_FONT
    ws3['A2'] = '任務：(一)3筆Follow-up歸位並標記完成 (二)2筆PLT Check口頭說機台並標記完成'
    kv_table(ws3, 4, [
        ('開始時間', ''), ('完成時間', ''), ('耗時(秒)', ''),
        ('使用模式', '□實體作業(應選) □誤用數位閱片'),
    ])
    ws3['A10'] = '(一) Follow-up 歸位'
    h2 = ['檢體', '標準位置', '口頭答案', '標記完成', '正確']
    for c, h in enumerate(h2, 1):
        ws3.cell(row=11, column=c, value=h)
    style_header_row(ws3, 11, len(h2))
    fu = [
        ['H5080706280', '090-2', '', '□是 □否', ''],
        ['H5080720696', '125-12', '', '□是 □否', ''],
        ['H5080721201', '208-6', '', '□是 □否', ''],
    ]
    for i, row in enumerate(fu, 12):
        for c, v in enumerate(row, 1):
            ws3.cell(row=i, column=c, value=v)
    apply_table_borders(ws3, 11, 14, 1, 5)
    ws3['A16'] = '(二) PLT Check 機台'
    for c, h in enumerate(h2, 1):
        ws3.cell(row=17, column=c, value=h)
    style_header_row(ws3, 17, len(h2))
    plt = [
        ['H5080720847', 'DI1', '', '□是 □否', ''],
        ['H5080720647', 'DI2', '', '□是 □否', ''],
    ]
    for i, row in enumerate(plt, 18):
        for c, v in enumerate(row, 1):
            ws3.cell(row=i, column=c, value=v)
    apply_table_borders(ws3, 17, 19, 1, 5)
    ws3['A21'] = 'SEQ — 情境二'
    ws3['A21'].font = SECTION_FONT
    kv_table(ws3, 22, [
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'),
        ('受試者理由', ''),
    ])
    set_col_widths(ws3, [16, 12, 18, 14, 8])

    # --- Sheet: 情境三 ---
    ws4 = wb.create_sheet('情境三')
    ws4['A1'] = '情境三：歷史數據與臨床資訊缺失'
    ws4['A1'].font = SECTION_FONT
    ws4['A2'] = (
        '任務：你是負責數位閱片/顯微鏡檢的醫檢師。正在數位閱片模式下，發現三筆AI分類警示檢體，'
        '需要確認病人歷史報告。檢體：H5080706286、H5080720647、H5080721401。'
        '在模擬介面下確認病患歷史報告，並口頭說明與本次報告的關聯。'
    )
    kv_table(ws4, 4, [
        ('開始時間', ''), ('完成時間', ''), ('3筆總耗時(秒)', ''),
    ])
    h3 = ['序', '檢體', '前次標準答案', '口頭答案摘要', '關聯合理', '耗時(秒)', '正確']
    for c, h in enumerate(h3, 1):
        ws4.cell(row=9, column=c, value=h)
    style_header_row(ws4, 9, len(h3))
    s3 = [
        ['1', 'H5080706286', 'Blast 1%', '', '□是 □否', '', ''],
        ['2', 'H5080720647', 'Blast 2%', '', '□是 □否', '', ''],
        ['3', 'H5080721401', 'WBC 10', '', '□是 □否', '', ''],
    ]
    for i, row in enumerate(s3, 10):
        for c, v in enumerate(row, 1):
            ws4.cell(row=i, column=c, value=v)
    apply_table_borders(ws4, 9, 12, 1, 7)
    ws4['A14'] = '關聯參考：6286前次已有Blast→追蹤；0647 2%→6%上升；1401 WBC 10→32.5急升'
    ws4['A16'] = 'SEQ — 情境三'
    ws4['A16'].font = SECTION_FONT
    kv_table(ws4, 17, [
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'),
        ('受試者理由', ''),
    ])
    set_col_widths(ws4, [5, 16, 14, 24, 12, 10, 8])

    # --- Sheet: 情境四五 ---
    ws5 = wb.create_sheet('情境四五')
    ws5['A1'] = '情境四～五：異常警示不足／留單警示不足'
    ws5['A1'].font = SECTION_FONT
    ws5['A2'] = '任務：數位閱片模式下依序完成10筆閱片與報告簽核'
    kv_table(ws5, 4, [
        ('開始時間', ''), ('結束時間', ''), ('10筆總耗時', ''),
    ])
    ws5['A9'] = '(一) 閱片時異常發現'
    ws5['A9'].font = Font(bold=True, size=11)
    h4a = ['序', '檢體', '應回報異常', '標準', '是否回報', '回報時機', '正確']
    for c, h in enumerate(h4a, 1):
        ws5.cell(row=10, column=c, value=h)
    style_header_row(ws5, 10, len(h4a))
    anom = [
        ['1', 'H5080721101', '否', '正常', '', '', 'N/A'],
        ['2', 'H5080721102', '否', '正常', '', '', 'N/A'],
        ['3', 'H5080720647', '是', 'Blast 6%真實異常', '', '□閱片中 □報告後 □未回報', ''],
        ['4', 'H5080721103', '是', 'Promyelocyte 3% Present', '', '', ''],
        ['5', 'H5080721104', '是', 'AI偽陽性：3顆Lymph→Blast 3%', '', '', ''],
        ['6', 'H5080721301', '是', 'AI偽陽性(Lymph→Blast 2%)', '', '', ''],
        ['7', 'H5080721105', '否', '正常', '', '', 'N/A'],
        ['8', 'H5080706286', '是', 'AI Blast 2%需確認', '', '', ''],
        ['9', 'H5080721106', '否', '正常', '', '', 'N/A'],
        ['10', 'H5080721107', '否', '正常', '', '', 'N/A'],
    ]
    for i, row in enumerate(anom, 11):
        for c, v in enumerate(row, 1):
            ws5.cell(row=i, column=c, value=v)
    apply_table_borders(ws5, 10, 20, 1, 7)
    ws5['A22'] = '(一)小計：異常正確回報 ____ / 5  漏報____  誤報____'
    ws5['A24'] = '(二) 留單判斷'
    ws5['A24'].font = Font(bold=True, size=11)
    h4b = ['序', '檢體', '應口頭留單', '橫幅(主持人)', '是否回報', '回報時機', '正確']
    for c, h in enumerate(h4b, 1):
        ws5.cell(row=25, column=c, value=h)
    style_header_row(ws5, 25, len(h4b))
    leave = [
        ['1', 'H5080721101', '否', '綠', '', '', 'N/A'],
        ['2', 'H5080721102', '否', '綠', '', '', 'N/A'],
        ['3', 'H5080720647', '是', '黃-延續', '', '', ''],
        ['4', 'H5080721103', '是', '紅-新發', '', '', ''],
        ['5', 'H5080721104', '視判斷', '紅*', '', '', ''],
        ['6', 'H5080721301', '視判斷', '紅*', '', '', ''],
        ['7', 'H5080721105', '否', '綠', '', '', 'N/A'],
        ['8', 'H5080706286', '是', '黃-延續', '', '', ''],
        ['9', 'H5080721106', '否', '綠', '', '', 'N/A'],
        ['10', 'H5080721107', '否', '綠', '', '', 'N/A'],
    ]
    for i, row in enumerate(leave, 26):
        for c, v in enumerate(row, 1):
            ws5.cell(row=i, column=c, value=v)
    apply_table_borders(ws5, 25, 35, 1, 7)
    ws5['A37'] = '(二)小計：留單正確 ____ / 3  1301: □偽陽性 □誤判留單 □未提及'
    ws5['A38'] = '異常口頭回報順序(非SEQ)：'
    ws5['A39'] = 'SEQ-情境四(異常)：分數 □1-7  理由：'
    ws5['A40'] = 'SEQ-情境五(留單)：分數 □1-7  理由：'
    ws5['A41'] = 'SEQ-合併(備選)：分數 □1-7  理由：'
    set_col_widths(ws5, [5, 16, 12, 22, 12, 22, 8])

    # --- Sheet: SEQ總表 ---
    ws6 = wb.create_sheet('SEQ總表')
    ws6['A1'] = '階段二 SEQ 總表'
    ws6['A1'].font = SECTION_FONT
    h6 = ['情境', 'SEQ(1-7)', '備註']
    for c, h in enumerate(h6, 1):
        ws6.cell(row=3, column=c, value=h)
    style_header_row(ws6, 3, len(h6))
    seq_rows = [
        ['情境一', '', ''],
        ['情境二', '', ''],
        ['情境三', '', ''],
        ['情境四～五', '', '□合併一題請註明'],
        ['平均(後算)', '', ''],
    ]
    for i, row in enumerate(seq_rows, 4):
        for c, v in enumerate(row, 1):
            ws6.cell(row=i, column=c, value=v)
    apply_table_borders(ws6, 3, 8, 1, 3)
    set_col_widths(ws6, [22, 12, 36])

    # --- Sheet: 階段三 SUS ---
    ws_sus = wb.create_sheet('階段三SUS')
    ws_sus['A1'] = '階段三：SUS 改善前後比較'
    ws_sus['A1'].font = SECTION_FONT
    ws_sus['A2'] = '施測順序：先填改良版(B)，再回想現行版(A)。量表1-5(非常不同意～非常同意)'
    sus_q = [
        '我認為我會想要經常使用這個系統',
        '我覺得這個系統不必要地複雜',
        '我認為這個系統很容易使用',
        '我認為我需要技術支援才能使用這個系統',
        '我發現這個系統的各種功能整合得很好',
        '我認為這個系統有太多不一致之處',
        '我想大多數人都能很快學會使用這個系統',
        '我覺得這個系統使用起來很笨拙',
        '我使用這個系統時感到很自信',
        '我需要先學很多東西才能開始使用這個系統',
    ]
    h_sus = ['題號', '題目', 'A現行版回想', 'B改良版']
    for c, h in enumerate(h_sus, 1):
        ws_sus.cell(row=4, column=c, value=h)
    style_header_row(ws_sus, 4, len(h_sus))
    for i, q in enumerate(sus_q, 1):
        r = 4 + i
        ws_sus.cell(row=r, column=1, value=i)
        ws_sus.cell(row=r, column=2, value=q)
        ws_sus.cell(row=r, column=3, value='')
        ws_sus.cell(row=r, column=4, value='')
    apply_table_borders(ws_sus, 4, 14, 1, 4)
    ws_sus['A16'] = 'SUS總分(0-100，後算)：A現行版____  B改良版____  差異(B-A)____'
    ws_sus['A17'] = '開放題：改良版最大改善？'
    ws_sus['A18'] = '開放題：仍最困擾？'
    set_col_widths(ws_sus, [6, 48, 14, 14])

    # --- Sheet: 質性觀察 ---
    ws7 = wb.create_sheet('質性觀察')
    ws7['A1'] = '附錄：質性觀察（選填）'
    ws7['A1'].font = SECTION_FONT
    kv_table(ws7, 3, [
        ('操作策略', ''), ('介面痛點', ''), ('與現行流程差異', ''), ('技術／中斷備註', ''),
    ])
    set_col_widths(ws7, [22, 58])

    wb.save(path)


def set_cell_font(cell, font_name='Microsoft JhengHei', size=11, bold=False):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            run.font.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_kv_table(doc, items):
    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'
    for i, (k, v) in enumerate(items):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for p in table.rows[i].cells[0].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    doc.add_paragraph()


def build_word(path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    t = doc.add_heading('林口長庚血球型態分類軟體', level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading('改善後成效調查 — 主持人觀察紀錄表', level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('用途：主持人／觀察員離線對照用（受試者不可見）｜全程約 40 分鐘')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('一、研究目的', level=2)
    doc.add_paragraph(
        '驗證改善後模擬介面，是否在五項高提及率痛點上，比現行軟體更能支援醫檢師完成血球型態閱片相關判斷與作業。'
    )

    doc.add_heading('二、整體流程與時間', level=2)
    add_table(doc, ['階段', '內容', '時間'], [
        ['階段一', '簡介與受試者同意', '3 min'],
        ['導覽', '原型介面導覽介紹', '5 min'],
        ['階段二', '五項情境操作', '20 min'],
        ['階段三', 'SUS 改善前後比較', '10 min'],
    ])

    doc.add_heading('階段一、簡介與受試者同意', level=2)
    doc.add_paragraph(
        '簡介：本研究請您聆聽簡介後，觀看介面導覽、完成 5 項模擬情境操作並填寫 SUS 問卷，全程約 40 分鐘；'
        '所使用為假資料模擬介面，不影響真實 LIS 或病患報告。報告僅以 P1～P5 編號呈現；參與自願可隨時退出。'
    )
    add_kv_table(doc, [
        ('測試日期', ''), ('受試者編號', ''),
        ('職級', '□ 醫檢師  □ 住院醫師  □ 其他：________'),
        ('年資', '□ <1年  □ 1-3年  □ 3-5年  □ >5年'),
        ('主持人', ''), ('觀察員', ''),
        ('受試者簽名', ''), ('說明人簽名', ''), ('同意時間', ''),
    ])

    doc.add_heading('階段二、五項情境操作評估', level=2)
    doc.add_paragraph('重點：任務成功率、完成時間、SEQ（每情境結束後立即施測）')
    doc.add_paragraph('SEQ問句：「整體而言，剛才這項任務對您來說有多困難或多容易？」（1-7）')

    doc.add_page_break()
    doc.add_heading('情境一：需推片確認資訊缺失', level=2)
    doc.add_paragraph('任務：從 10 筆中找出所有仍需「需拉片確認(Follow-up)」的檢體，口頭說明編號與判斷依據。')
    add_kv_table(doc, [
        ('開始時間', ''), ('完成時間', ''), ('耗時(秒)', ''),
        ('是否說明判斷依據', '□ 是  □ 否'),
        ('是否誤選 H5080706301', '□ 是  □ 否'), ('正確數', '____ / 3'),
    ])
    add_table(doc, ['檢體編號', '是否應選', '判斷依據', '受試者是否提及', '□✓ □✗'], [
        ['H5080706280', '應選', 'Follow-up待辦；歸位090-2', '', ''],
        ['H5080720696', '應選', 'Follow-up待辦；歸位125-12', '', ''],
        ['H5080721201', '應選', 'AI+Follow-up雙旗標；歸位208-6', '', ''],
        ['H5080706301', '不可選', '陷阱：Follow-up已完成', '', ''],
        ['其餘6筆', '不可選', '非待辦Follow-up', '', ''],
    ])
    doc.add_heading('SEQ — 情境一', level=3)
    add_kv_table(doc, [
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'), ('受試者理由（一句話）', ''),
    ])

    doc.add_page_break()
    doc.add_heading('情境二：檢體實體位置／機台位置缺失', level=2)
    doc.add_paragraph('任務：(一)3筆Follow-up歸位並標記完成 (二)2筆PLT Check口頭說機台並標記完成')
    add_kv_table(doc, [
        ('開始時間', ''), ('完成時間', ''), ('耗時(秒)', ''),
        ('使用模式', '□ 實體作業(應選)  □ 誤用數位閱片'),
    ])
    doc.add_paragraph('(一) Follow-up 歸位 — ____ / 3')
    add_table(doc, ['檢體', '標準位置', '口頭答案', '標記完成', '□✓ □✗'], [
        ['H5080706280', '090-2', '', '□ 是 □ 否', ''],
        ['H5080720696', '125-12', '', '□ 是 □ 否', ''],
        ['H5080721201', '208-6', '', '□ 是 □ 否', ''],
    ])
    doc.add_paragraph('(二) PLT Check 機台 — ____ / 2')
    add_table(doc, ['檢體', '標準機台', '口頭答案', '標記完成', '□✓ □✗'], [
        ['H5080720847', 'DI1', '', '□ 是 □ 否', ''],
        ['H5080720647', 'DI2', '', '□ 是 □ 否', ''],
    ])
    doc.add_heading('SEQ — 情境二', level=3)
    add_kv_table(doc, [
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'), ('受試者理由（一句話）', ''),
    ])

    doc.add_page_break()
    doc.add_heading('情境三：歷史數據與臨床資訊缺失', level=2)
    doc.add_paragraph(
        '任務：你是負責數位閱片/顯微鏡檢的醫檢師。正在數位閱片模式下，發現三筆AI分類警示檢體，'
        '需要確認病人歷史報告。檢體編號：H5080706286、H5080720647、H5080721401。'
        '在模擬介面下確認病患歷史報告，並口頭說明與本次報告的關聯。'
    )
    add_kv_table(doc, [('開始時間', ''), ('完成時間', ''), ('3筆總耗時(秒)', '')])
    add_table(doc, ['序', '檢體', '前次標準答案', '口頭答案摘要', '關聯合理', '耗時(秒)', '□✓ □✗'], [
        ['1', 'H5080706286', 'Blast 1%', '', '□ 是 □ 否', '', ''],
        ['2', 'H5080720647', 'Blast 2%', '', '□ 是 □ 否', '', ''],
        ['3', 'H5080721401', 'WBC 10', '', '□ 是 □ 否', '', ''],
    ])
    doc.add_paragraph('關聯參考：6286前次已有Blast→追蹤；0647 2%→6%上升；1401 WBC 10→32.5急升')
    doc.add_heading('SEQ — 情境三', level=3)
    add_kv_table(doc, [
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'), ('受試者理由（一句話）', ''),
    ])

    doc.add_page_break()
    doc.add_heading('情境四～五：異常警示不足／留單警示不足', level=2)
    doc.add_paragraph('任務：數位閱片模式下依序完成 10 筆閱片與報告簽核。(一)異常發現口頭告知 (二)需留單口頭告知')
    add_kv_table(doc, [('開始時間', ''), ('結束時間', ''), ('10筆總耗時', '')])
    doc.add_paragraph('(一) 閱片時異常發現')
    add_table(doc, ['序', '檢體', '應回報', '標準', '是否回報', '回報時機', '□✓ □✗ □N/A'], [
        ['1', 'H5080721101', '否', '正常', '', '', 'N/A'],
        ['2', 'H5080721102', '否', '正常', '', '', 'N/A'],
        ['3', 'H5080720647', '是', 'Blast 6%真實異常', '', '□閱片中 □報告後 □未回報', ''],
        ['4', 'H5080721103', '是', 'Promyelocyte 3%', '', '', ''],
        ['5', 'H5080721104', '是', 'AI偽陽性：3顆Lymph→Blast', '', '', ''],
        ['6', 'H5080721301', '是', 'AI偽陽性', '', '', ''],
        ['7', 'H5080721105', '否', '正常', '', '', 'N/A'],
        ['8', 'H5080706286', '是', 'AI Blast 2%需確認', '', '', ''],
        ['9', 'H5080721106', '否', '正常', '', '', 'N/A'],
        ['10', 'H5080721107', '否', '正常', '', '', 'N/A'],
    ])
    doc.add_paragraph('(一)小計：異常正確回報 ____ / 5  ｜  漏報 ____  ｜  誤報 ____')
    doc.add_paragraph('(二) 留單判斷')
    add_table(doc, ['序', '檢體', '應留單', '橫幅(主持人)', '是否回報', '回報時機', '□✓ □✗ □N/A'], [
        ['1', 'H5080721101', '否', '綠', '', '', 'N/A'],
        ['2', 'H5080721102', '否', '綠', '', '', 'N/A'],
        ['3', 'H5080720647', '是', '黃-延續', '', '', ''],
        ['4', 'H5080721103', '是', '紅-新發', '', '', ''],
        ['5', 'H5080721104', '視判斷', '紅*', '', '', ''],
        ['6', 'H5080721301', '視判斷', '紅*', '', '', ''],
        ['7', 'H5080721105', '否', '綠', '', '', 'N/A'],
        ['8', 'H5080706286', '是', '黃-延續', '', '', ''],
        ['9', 'H5080721106', '否', '綠', '', '', 'N/A'],
        ['10', 'H5080721107', '否', '綠', '', '', 'N/A'],
    ])
    doc.add_paragraph('(二)小計：留單正確 ____ / 3  ｜  1301：□偽陽性 □誤判留單 □未提及')
    doc.add_paragraph('異常口頭回報順序(非SEQ)：________________________________')
    doc.add_heading('SEQ — 情境四（異常發現）', level=3)
    add_kv_table(doc, [
        ('問句', '「整體而言，剛才連續閱片時發現並回報異常，對您來說有多困難或多容易？」'),
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'), ('受試者理由', ''),
    ])
    doc.add_heading('SEQ — 情境五（留單判斷）', level=3)
    add_kv_table(doc, [
        ('問句', '「整體而言，剛才在閱片過程中判斷是否需要留單，對您來說有多困難或多容易？」'),
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'), ('受試者理由', ''),
    ])
    doc.add_heading('SEQ — 情境四～五合併（備選）', level=3)
    add_kv_table(doc, [
        ('問句', '「整體而言，剛才完成10筆閱片與簽核(含異常發現與留單判斷)，對您來說有多困難或多容易？」'),
        ('SEQ分數(1-7)', '□1 □2 □3 □4 □5 □6 □7'), ('受試者理由', ''),
    ])

    doc.add_page_break()
    doc.add_heading('階段三、SUS 改善前後比較', level=2)
    doc.add_paragraph('施測順序：先填改良版(B)，再回想現行版(A)。量表 1-5（非常不同意～非常同意）')
    add_table(doc, ['題號', '題目', 'A現行版', 'B改良版'], [
        ['1', '我認為我會想要經常使用這個系統', '', ''],
        ['2', '我覺得這個系統不必要地複雜', '', ''],
        ['3', '我認為這個系統很容易使用', '', ''],
        ['4', '我認為我需要技術支援才能使用這個系統', '', ''],
        ['5', '我發現這個系統的各種功能整合得很好', '', ''],
        ['6', '我認為這個系統有太多不一致之處', '', ''],
        ['7', '我想大多數人都能很快學會使用這個系統', '', ''],
        ['8', '我覺得這個系統使用起來很笨拙', '', ''],
        ['9', '我使用這個系統時感到很自信', '', ''],
        ['10', '我需要先學很多東西才能開始使用這個系統', '', ''],
    ])
    add_kv_table(doc, [
        ('SUS總分A(現行版)', ''), ('SUS總分B(改良版)', ''), ('差異(B-A)', ''),
        ('改良版最大改善', ''), ('仍最困擾', ''),
    ])

    doc.add_heading('附錄：質性觀察（選填）', level=2)
    add_kv_table(doc, [
        ('操作策略', ''), ('介面痛點', ''), ('與現行流程差異', ''), ('技術／中斷備註', ''),
    ])

    doc.save(path)


if __name__ == '__main__':
    xlsx = os.path.join(OUT_DIR, '主持人觀察紀錄表.xlsx')
    docx_path = os.path.join(OUT_DIR, '主持人觀察紀錄表.docx')
    build_excel(xlsx)
    build_word(docx_path)
    print('Generated:')
    print(' ', xlsx)
    print(' ', docx_path)
