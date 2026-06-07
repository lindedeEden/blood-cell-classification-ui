# -*- coding: utf-8 -*-
"""產生「血球分類軟體介面專案」簡報說明 Word 文件。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "血球分類軟體介面專案_簡報說明文件.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, 16 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, 10, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    set_run_font(run, 10)
    doc.add_paragraph()


def build():
    doc = Document()
    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("血球型態分類軟體介面專案\n簡報說明文件")
    set_run_font(r, 22, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("林口長庚醫院檢驗醫學科血液組｜前端介面 Demo\n（供檢驗科主任、AI 中心研究員簡報用）")
    set_run_font(r2, 12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("文件版本：2026-05-30｜性質：操作與設計說明（非正式 LIS 規格書）")
    set_run_font(r3, 10, color=(100, 100, 100))

    doc.add_page_break()

    # 0. 如何使用本文件
    add_title(doc, "如何使用本文件", 1)
    add_para(doc, "本文件協助您向兩類聽眾介紹同一套 Demo，但切入角度不同：")
    add_table(doc,
              ["聽眾", "關注重點", "建議閱讀章節"],
              [
                  ["檢驗科主任", "臨床流程是否合理、能否減輕醫檢師負擔、留單與簽核是否可控", "第 1～4 章、第 6 章"],
                  ["AI 中心研究員", "AI 輸出如何呈現、人如何覆寫、留單判定、狀態與 API 接軌點", "第 1、3、5、7 章"],
                  ["兩者共同", "現場 Demo 怎麼走、示範檢體怎麼選", "第 6 章"],
              ])

    add_title(doc, "建議簡報時間配置（約 40～50 分鐘）", 2)
    add_table(doc,
              ["時間", "內容", "目的"],
              [
                  ["5 分", "開場：為什麼要做這個介面", "對齊臨床痛點與 AI 落地需求"],
                  ["8 分", "整體流程與雙軌分工（數位／實體）", "讓主任先建立全局觀"],
                  ["10 分", "Live Demo：檢體管理 → 閱片 → 報告", "用畫面說話"],
                  ["8 分", "AI 整合與人機協作設計", "回應研究員技術問題"],
                  ["5 分", "留單規則、風險橫幅、簽核邏輯", "展現臨床安全設計"],
                  ["5 分", "後續接軌與 Q&A", "收斂期待"],
              ])

    doc.add_page_break()

    # 1. 專案定位
    add_title(doc, "第 1 章　專案定位：這是什麼、不是什麼", 1)

    add_title(doc, "1.1 這是什麼", 2)
    add_bullets(doc, [
        "一套「血球型態 AI 輔助分類」的前端操作介面 Demo，模擬醫檢師日常路徑。",
        "完整路徑：登入 → 檢體管理 → 影像閱片與細胞編輯 → 報告核發。",
        "內建模擬資料（database.js）與真實血球範例影像，可在瀏覽器直接操作。",
        "目的：在接 LIS／AI 後端之前，先與臨床、AI 團隊對齊「畫面與流程」。",
    ])

    add_title(doc, "1.2 這不是什麼（請先講清楚，避免期待落差）", 2)
    add_bullets(doc, [
        "不是已上線的 LIS 模組，目前不會真的寫入檢驗報告系統。",
        "不是 AI 模型本身；模型推論結果以「AI 欄位」呈現，可替換為 API 回傳。",
        "不是完整資安／權限系統；登入為 Demo 帳密（admin/admin）。",
        "部分狀態存在瀏覽器本機（localStorage），重新登入會重置部分示範狀態。",
    ])

    add_title(doc, "1.3 解決的臨床問題（給主任的開場白）", 2)
    add_bullets(doc, [
        "數位顯微影像＋AI 分類後，醫檢師仍需：看細胞、改分類、對照前次、判斷是否留片。",
        "數位閱片與實體拉片／PLT 確認常由不同人、不同時段處理，需要清楚待辦分流。",
        "AI 可能有誤報；介面需支援「確認、覆寫、改走實體鏡檢」而不打亂流程。",
        "留單標準需可設定，並區分「新發異常」與「延續性異常」，避免重複留片。",
    ])

    doc.add_page_break()

    # 2. 使用者與分工
    add_title(doc, "第 2 章　使用者角色與工作分工", 1)
    add_table(doc,
              ["角色", "主要使用畫面", "負責工作"],
              [
                  ["鏡檢醫檢師", "數位閱片模式、影像閱片頁", "看細胞、改分類、數位簽核或轉實體拉片"],
                  ["報告／核發醫檢師", "實體作業模式、報告核發", "PLT 確認、拉片完成、LIS 報告核發（Demo 模擬）"],
                  ["組長／主任", "檢體管理總覽", "掌握待辦、時效、異常分布（Demo 為單機示意）"],
                  ["AI 工程師", "資料欄位、狀態 API", "對接 metrics、細胞分類、信心度、模型版本"],
              ])

    add_title(doc, "2.1 雙軌流程（核心概念，建議用簡報投影片放一張圖）", 2)
    add_para(doc, "每筆檢體有兩條可並行的流程線：")
    add_bullets(doc, [
        "【數位流程】數位閱片、AI 分類警示 — 在螢幕上完成判讀與簽核。",
        "【實體流程】需拉片確認、PLT Check — 到實驗室看玻片或確認儀器異常。",
        "整件結案 = 兩邊「該做的都做完了」。",
        "若檢體沒有某類標籤（例如只有 AI 警示、沒有 PLT），該邊視為無須處理。",
    ])

    doc.add_page_break()

    # 3. 畫面導覽
    add_title(doc, "第 3 章　四大畫面導覽", 1)

    add_title(doc, "3.1 登入頁（index.html）", 2)
    add_bullets(doc, [
        "Demo 帳密：admin/admin 或 user/user。",
        "登入後記錄操作者帳號，供檢體「整件完成」時寫入編輯人員欄。",
        "重新登入會清除部分本機示範狀態，方便從頭 Demo。",
    ])

    add_title(doc, "3.2 檢體管理（檢體管理.html）— 工作台首頁", 2)
    add_bullets(doc, [
        "【數位閱片模式】：列出尚待完成數位工作的檢體。",
        "【實體作業模式】：列出需拉片、PLT 等實體待辦。",
        "右側總覽：檢體資訊、流式計數／AI／前次報告對照表。",
        "狀態膠囊：一眼看出待辦類型；完成後變綠色打勾。",
        "系統設定：可調整留單門檻（與閱片、報告共用）。",
        "進入閱片／唯讀檢視；數位已簽核結案者只能唯讀。",
    ])

    add_title(doc, "3.3 影像閱片與細胞編輯（影像檢視與細胞編輯.html）", 2)
    add_bullets(doc, [
        "左側：檢體資訊、分析表、狀態膠囊、上一筆／下一筆。",
        "主區：依細胞類別分群；異常群組（如 Blast）置頂並紅色強調。",
        "操作：點選、多選、拖曳改分類；支援單手模式與影像縮放。",
        "進度：須全部細胞看過且無 Unidentified 才能開報告。",
        "儲存並核發報告：開啟報告核發視窗（iframe）。",
        "使用真實血球範例影像（依類別對應 sample 圖）。",
    ])

    add_title(doc, "3.4 報告核發（報告核發.html）", 2)
    add_bullets(doc, [
        "WBC 分類表：流式計數｜AI｜人員編輯｜前次報告。",
        "流式計數：成熟細胞有值；未成熟／異常細胞一律「-」（儀器無法提供）。",
        "風險橫幅：綠（正常）／黃（延續性異常）／紅（新發留單）。",
        "簽核按鈕：完成數位流程；紅色或待拉片時預設上鎖，可解鎖強制簽核。",
        "改為人工鏡檢：紅色新發留單時，可轉交實體拉片（數位標完成但未簽核結案）。",
    ])

    doc.add_page_break()

    # 4. 狀態膠囊
    add_title(doc, "第 4 章　狀態膠囊與臨床意義（給主任）", 1)
    add_table(doc,
              ["膠囊", "中文", "代表意思", "誰處理", "完成後"],
              [
                  ["Digital Review", "數位閱片", "需要在螢幕上完成細胞檢視與判讀", "鏡檢醫檢師", "綠勾；簽核結案後唯讀"],
                  ["AI Alert", "AI 分類警示", "AI 發現需人工確認的分類結果", "鏡檢醫檢師", "確認後綠勾"],
                  ["Follow-up", "需拉片確認", "需到實體玻片再確認（含追蹤、轉實體）", "實體醫檢師", "拉片完成後綠勾"],
                  ["PLT Check", "血小板確認", "PLT 與前次或臨床不符，需鏡檢", "實體醫檢師", "確認後綠勾"],
              ])

    add_title(doc, "4.1 特殊組合規則", 2)
    add_bullets(doc, [
        "AI + 需拉片 同時存在：只在「實體作業」清單出現（數位清單排除）。",
        "數位已完成、只剩 PLT：從數位清單消失，留實體清單。",
        "整件已完成：時效歸零、可篩選 Verified；編輯人員欄寫入最後操作者。",
    ])

    doc.add_page_break()

    # 5. AI 整合
    add_title(doc, "第 5 章　AI 整合與人機協作（給 AI 中心研究員）", 1)

    add_title(doc, "5.1 資料欄位設計", 2)
    add_table(doc,
              ["欄位", "來源", "用途"],
              [
                  ["metrics.*", "AI 模型／推論 API", "報告「AI」欄、風險判定基準"],
                  ["flowCyt.*", "流式細胞儀", "報告「流式計數」欄；異常細胞固定「-」"],
                  ["editedMetrics.*", "醫檢師人工修改", "報告「人員編輯」欄；覆寫 AI 參與留單判定"],
                  ["prevReport.*", "LIS 前次報告", "延續性 vs 新發留單、前次對照"],
                  ["cells[]", "AI 偵測＋分類", "閱片頁細胞群組、thumbnail、類別標籤"],
              ])

    add_title(doc, "5.2 人機協作的三層判斷", 2)
    add_bullets(doc, [
        "第一層 — AI 自動分類：提供各類細胞比例與細胞影像分群。",
        "第二層 — 人工覆寫：醫檢師可拖曳改類別；editedMetrics 與 AI 不同時，留單以人員編輯為準。",
        "第三層 — 規則引擎：LEAVE_THRESHOLDS 依林口留單標準判定；比對 prevReport 區分新發／延續。",
    ])

    add_title(doc, "5.3 AI 輸出與 UI 回饋", 2)
    add_bullets(doc, [
        "AI Alert 膠囊：模型結果達門檻或需確認時標記。",
        "異常細胞列紅色強調（列表、側欄、報告表一致）。",
        "誤報情境：人員改為正常 → 綠色橫幅 → 可直接簽核確認 AI（不需轉實體）。",
        "漏報／新發：人員改出異常或 AI 新發留單 → 紅色橫幅 → 鎖簽核或改走實體。",
    ])

    add_title(doc, "5.4 後端接軌建議（研究員可能會問）", 2)
    add_table(doc,
              ["接軌點", "現況", "建議 API"],
              [
                  ["檢體清單", "database.js", "GET /specimens?date=&status=&mode="],
                  ["細胞影像與分類", "前端 mock + sample 圖", "GET /specimens/{id}/cells"],
                  ["AI metrics", "metrics 物件", "POST 推論結果或 GET 已存結果"],
                  ["簽核", "postMessage reportVerified", "POST /specimens/{id}/sign-off"],
                  ["狀態更新", "localStorage 覆寫", "PATCH /specimens/{id}/workflow"],
              ])

    add_title(doc, "5.5 可延伸的 AI 研究議題", 2)
    add_bullets(doc, [
        "模型信心度（per-cell / per-class）如何顯示在 UI？",
        "Active learning：醫檢師改分類是否回饋訓練？editedMetrics 即為標註來源。",
        "人機一致性指標：AI vs 人員編輯差異率、誤報率統計。",
        "多模型版本 A/B：metrics 可加 modelVersion 欄位。",
    ])

    doc.add_page_break()

    # 6. Demo 腳本
    add_title(doc, "第 6 章　現場 Demo 建議腳本（約 10 分鐘）", 1)
    add_para(doc, "操作前準備：Chrome/Edge 開啟 index.html → 登入 admin → 日期設 2025-08-07。")

    add_title(doc, "路線 A：一般數位閱片（給主任看「正常流程」）", 2)
    add_bullets(doc, [
        "1. 檢體管理 → 數位閱片模式 → 選 H5080721101（楊建志，僅數位閱片）。",
        "2. 進入閱片 → 展示細胞分群、進度條、縮放。",
        "3. 全部看過 → 儲存並核發報告 → 綠色橫幅 → 確認並簽核。",
        "4. 回列表：數位閱片綠勾、可唯讀檢視。",
    ])

    add_title(doc, "路線 B：AI 誤報排除（給 AI 研究員看「人覆寫 AI」）", 2)
    add_bullets(doc, [
        "1. 選 H5080706286（李俊傑，AI Alert：Blast 2%）。",
        "2. 閱片確認其實沒有 Blast → 或不改直接看報告。",
        "3. 報告綠色橫幅 → 簽核 → AI 確認、補數位閱片綠勾。",
        "重點：AI 警示不一定代表要留片，人員可確認後簽核。",
    ])

    add_title(doc, "路線 C：新發留單 → 改走實體（核心安全設計）", 2)
    add_bullets(doc, [
        "1. 選 H5080721101 → 閱片時人工增加 Blast 分類。",
        "2. 報告紅色橫幅 → 簽核上鎖 → 出現「改為人工鏡檢」。",
        "3. 點改為人工鏡檢 → 數位綠勾（非唯讀）+ 需拉片待辦 → 回實體清單。",
        "4.（可選）解鎖強制簽核路線：展示「堅持用數位結果結案」的替代流程。",
    ])

    add_title(doc, "路線 D：實體追蹤（前次異常、本次正常）", 2)
    add_bullets(doc, [
        "1. 實體作業模式 → H5080706280（吳佩琪，需拉片確認）。",
        "2. 右側表：前次 Eos 25%、本次正常；說明追蹤案例設計。",
        "3. 實體完成拉片 → 點需拉片膠囊或報告「已拉片完成」。",
    ])

    add_title(doc, "路線 E：AI + 需拉片雙旗標（最複雜案例）", 2)
    add_bullets(doc, [
        "1. 實體模式 → H5080721201（蔡宗翰，AI + 需拉片）。",
        "2. 說明：只在實體清單；流式 Blast 為「-」，AI Blast 7%。",
        "3. 拉片完成時 AI 一併確認。",
    ])

    doc.add_page_break()

    # 7. 報告與留單
    add_title(doc, "第 7 章　報告核發、留單與簽核邏輯", 1)

    add_title(doc, "7.1 風險橫幅三色", 2)
    add_table(doc,
              ["顏色", "條件", "簽核", "臨床意義"],
              [
                  ["綠色", "未達留單門檻", "可直接簽核", "可正常發報告"],
                  ["黃色", "達門檻但前次已有", "可直接簽核", "延續性異常，不重複留片"],
                  ["紅色", "新發達門檻", "上鎖", "需確認是否留片／改實體"],
              ])

    add_title(doc, "7.2 報告頁按鈕決策", 2)
    add_table(doc,
              ["使用者動作", "結果"],
              [
                  ["確認並簽核（綠/黃）", "數位簽核結案 → 唯讀"],
                  ["改為人工鏡檢（紅色）", "數位綠勾（可編輯）+ 需拉片待辦"],
                  ["解鎖 → 確認並簽核", "數位簽核結案 → 唯讀；不新增需拉片"],
                  ["已拉片完成", "需拉片綠勾 → 可簽核"],
              ])

    add_title(doc, "7.3 留單門檻（節錄林口標準）", 2)
    add_bullets(doc, [
        "WBC ≧ 30,000/μL；Lymphocyte 成人 ≧ 60%；Eosinophil ≧ 20%。",
        "Blast / Promyelocyte：Present 即留；Myelocyte ≧ 5%；Metamyelocyte ≧ 10%。",
        "完整標準見專案「附件一 林口長庚醫院留單標準.txt」；系統設定可調整 Demo 門檻。",
    ])

    doc.add_page_break()

    # 8. 技術附錄
    add_title(doc, "第 8 章　技術附錄（給研究員／資訊人員）", 1)

    add_title(doc, "8.1 專案結構", 2)
    add_bullets(doc, [
        "index.html、檢體管理.html、影像檢視與細胞編輯.html、報告核發.html",
        "assets/data/database.js — 模擬檢體與 metrics",
        "assets/js/common.js — 流程狀態、留單、導頁、持久化",
        "assets/js/image-review.js — 閱片、細胞編輯、報告 iframe",
        "assets/js/report-issue.js — 報告表、風險橫幅、簽核",
    ])

    add_title(doc, "8.2 workflowDone 狀態（核心狀態機）", 2)
    add_table(doc,
              ["旗標", "意義"],
              [
                  ["digitalReview", "數位這段已處理 → 膠囊綠勾"],
                  ["digitalReviewSignedOff", "已確認並簽核 → 閱片唯讀"],
                  ["aiAlertConfirmed", "AI 警示已確認"],
                  ["entityStatusDone.Follow-up", "需拉片已完成"],
                  ["entityStatusDone.PLT Check", "PLT 確認已完成"],
              ])

    add_title(doc, "8.3 本機持久化（Demo 限制）", 2)
    add_bullets(doc, [
        "blood-morphology-specimen-status：檢體狀態覆寫",
        "editedMetrics:{檢體ID}：人員編輯的比例",
        "重新登入清除狀態覆寫；接後端時改為 API 同步",
    ])

    doc.add_page_break()

    # 9. Q&A
    add_title(doc, "第 9 章　預期提問與建議回答", 1)
    add_table(doc,
              ["可能的問題", "建議回答"],
              [
                  ["這什麼時候能上線？", "目前是流程與介面定稿 Demo；上線需接 LIS、資安、AI API，時程另訂。"],
                  ["AI 錯了怎麼辦？", "醫檢師可改分類；報告以人員編輯為準；可確認 AI 或轉實體拉片。"],
                  ["會不會重複留片？", "系統比對前次報告，延續性異常顯示黃色、不重複留片。"],
                  ["數位與實體誰做？", "清單分模式；雙旗標案例在實體模式統一處理。"],
                  ["模型輸出格式？", "metrics 物件 + cells 陣列；詳見第 5 章，可對接 REST API。"],
                  ["和現有 LIS 關係？", "本 Demo 獨立運行；簽核、狀態、報告皆預留 API 接點。"],
              ])

    add_title(doc, "第 10 章　簡報收尾建議", 1)
    add_bullets(doc, [
        "向主任強調：這套介面把「AI 建議、人工判讀、留片規則、簽核責任」分開且可追蹤。",
        "向 AI 研究員強調：UI 已預留 AI 輸入／人工標註／規則引擎三層，便於評估模型臨床可用性。",
        "下一步可討論：優先接軌項目（清單 API、推論 API、簽核 API）、Pilot 檢體量、評估指標。",
        "邀請現場操作 1～2 條 Demo 路線，比純投影片更有說服力。",
    ])

    doc.save(OUTPUT)
    print(f"已產生：{OUTPUT}")


if __name__ == "__main__":
    build()
