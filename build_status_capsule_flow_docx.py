# -*- coding: utf-8 -*-
"""產生「狀態膠囊流程說明」Word 文件（含流程圖與文字說明）。"""
from __future__ import annotations

import textwrap
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "_status_capsule_doc_assets"
OUTPUT = ROOT / "狀態膠囊流程說明.docx"
CAPSULE_STYLE_IMG = ROOT / "狀態膠囊樣式對照圖.png"


def setup_matplotlib_font():
    plt.rcParams["font.sans-serif"] = ["Microsoft JhengHei", "Microsoft YaHei", "SimHei", "sans-serif"]
    plt.rcParams["axes.unicode_minus"] = False


def set_run_font(run, size=11, bold=False, color=None, name="Microsoft JhengHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, 18 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, 10, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    set_run_font(run, 10)
    doc.add_paragraph()


def add_image(doc, path: Path, caption: str, width_in=6.2):
    if not path.exists():
        add_para(doc, f"（圖片缺失：{path.name}）", size=10)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, 10, color=(100, 100, 100))
    cap.paragraph_format.space_after = Pt(12)


class FlowChart:
    """以 matplotlib 繪製簡化流程圖。"""

    def __init__(self, title: str, figsize=(10, 7)):
        self.title = title
        self.fig, self.ax = plt.subplots(figsize=figsize)
        self.ax.set_xlim(0, 10)
        self.ax.set_ylim(0, 10)
        self.ax.axis("off")
        self.ax.set_title(title, fontsize=14, fontweight="bold", pad=12)

    def box(self, x, y, w, h, text, face="#eef2ff", edge="#4338ca", fontsize=9):
        wrap_w = max(8, int(w * 3.2))
        lines = textwrap.wrap(text, width=wrap_w) or [text]
        display = "\n".join(lines)
        patch = FancyBboxPatch(
            (x - w / 2, y - h / 2),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.08",
            linewidth=1.2,
            edgecolor=edge,
            facecolor=face,
        )
        self.ax.add_patch(patch)
        self.ax.text(x, y, display, ha="center", va="center", fontsize=fontsize)

    def diamond(self, x, y, size, text, face="#fff7ed", edge="#c2410c", fontsize=9):
        s = size / 2
        pts = [(x, y + s), (x + s, y), (x, y - s), (x - s, y), (x, y + s)]
        xs, ys = zip(*pts)
        self.ax.fill(xs, ys, facecolor=face, edgecolor=edge, linewidth=1.2)
        self.ax.text(x, y, textwrap.fill(text, width=10), ha="center", va="center", fontsize=fontsize)

    def arrow(self, x1, y1, x2, y2, label=""):
        arr = FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=1.1,
            color="#475569",
            shrinkA=4,
            shrinkB=4,
        )
        self.ax.add_patch(arr)
        if label:
            self.ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.15, label, ha="center", va="bottom", fontsize=8, color="#334155")

    def note(self, x, y, text, fontsize=8, color="#64748b"):
        self.ax.text(x, y, text, ha="center", va="center", fontsize=fontsize, color=color, style="italic")

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.fig.tight_layout()
        self.fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
        plt.close(self.fig)


def chart_overview(path: Path):
    fc = FlowChart("圖 1　系統總覽：雙流程與整體結案", figsize=(11, 8))
    fc.box(2.5, 8.6, 3.8, 0.9, "數位閱片模式\n待辦：DR 未完成 / AI 待確認", "#ede9fe", "#6d28d9")
    fc.box(7.5, 8.6, 3.8, 0.9, "實體作業模式\n待辦：PLT / Follow-up / AI+FU", "#dbeafe", "#1d4ed8")
    fc.box(2.5, 6.8, 3.2, 0.8, "[紫] 數位閱片 待辦", "#f3e8ff", "#7c3aed")
    fc.box(2.5, 5.5, 3.2, 0.8, "影像檢視 → 報告核發", "#f8fafc", "#64748b")
    fc.box(2.5, 4.2, 3.2, 0.8, "[完成] 數位閱片", "#dcfce7", "#15803d")
    fc.box(7.5, 6.8, 3.2, 0.8, "[藍] PLT / [紅] Follow-up", "#eff6ff", "#2563eb")
    fc.box(7.5, 5.5, 3.2, 0.8, "列表點膠囊 / 已拉片完成", "#f8fafc", "#64748b")
    fc.box(7.5, 4.2, 3.2, 0.8, "[完成] 實體膠囊", "#dcfce7", "#15803d")
    fc.box(5.0, 2.5, 4.5, 1.0, "整體流程完成\n數位邊 + 實體邊皆通過 → Verified 篩選", "#ecfdf5", "#047857", 10)
    fc.arrow(2.5, 8.1, 2.5, 7.25)
    fc.arrow(2.5, 6.35, 2.5, 5.95)
    fc.arrow(2.5, 5.05, 2.5, 4.65)
    fc.arrow(7.5, 8.1, 7.5, 7.25)
    fc.arrow(7.5, 6.35, 7.5, 5.95)
    fc.arrow(7.5, 5.05, 7.5, 4.65)
    fc.arrow(2.5, 3.75, 4.0, 3.0, "合流")
    fc.arrow(7.5, 3.75, 6.0, 3.0, "合流")
    fc.note(5.0, 1.2, "任一邊「無須處理」視為該邊已通過（如無 DR 膠囊、無實體膠囊）")
    fc.save(path)


def chart_digital_review(path: Path):
    fc = FlowChart("圖 2　數位閱片（Digital Review）生命週期", figsize=(10, 9))
    fc.box(5, 8.5, 4.2, 0.8, "[紫] 待辦：digitalReview = false", "#f3e8ff", "#7c3aed")
    fc.box(5, 7.2, 3.6, 0.7, "進入閱片（可編輯）", "#f8fafc", "#64748b")
    fc.box(5, 6.0, 3.6, 0.7, "儲存並核發報告", "#f8fafc", "#64748b")
    fc.diamond(5, 4.7, 1.6, "報告頁動作？")
    fc.box(2.0, 3.0, 3.4, 1.0, "綠/黃：確認並簽核\n[完成] + 簽核結案\nsignedOff = true", "#dcfce7", "#15803d", 8)
    fc.box(5.0, 3.0, 3.4, 1.0, "紅：改為人工鏡檢\n[完成]（交接）\n+ [紅] 需拉片確認", "#fef3c7", "#b45309", 8)
    fc.box(8.0, 3.0, 3.0, 1.0, "紅：開鎖強制簽核\n[完成] + 結案", "#fee2e2", "#b91c1c", 8)
    fc.box(2.0, 1.2, 3.2, 0.9, "唯讀：已簽核結案", "#fff7ed", "#9a3412", 8)
    fc.box(5.0, 1.2, 3.2, 0.9, "唯讀：交接快照\n→ 實體作業拉片", "#fff7ed", "#9a3412", 8)
    fc.box(8.0, 1.2, 3.0, 0.8, "右鍵退回 DR\n可再編輯", "#e0f2fe", "#0369a1", 8)
    fc.arrow(5, 8.05, 5, 7.6)
    fc.arrow(5, 6.8, 5, 6.4)
    fc.arrow(5, 5.6, 5, 5.45)
    fc.arrow(4.2, 4.0, 2.0, 3.55, "綠/黃")
    fc.arrow(5, 4.0, 5, 3.55, "改人工")
    fc.arrow(5.8, 4.0, 8.0, 3.55, "開鎖")
    fc.arrow(2.0, 2.45, 2.0, 1.7)
    fc.arrow(5.0, 2.45, 5.0, 1.7)
    fc.arrow(8.0, 2.45, 8.0, 1.65)
    fc.save(path)


def chart_ai_alert(path: Path):
    fc = FlowChart("圖 3　AI 分類警示（AI Alert）生命週期", figsize=(10, 7))
    fc.box(5, 8.2, 4.0, 0.8, "[橘] 待辦：aiAlertConfirmed = false", "#ffedd5", "#c2410c")
    fc.diamond(5, 6.7, 1.7, "檢體組合？")
    fc.box(2.3, 5.0, 3.4, 0.9, "僅 AI Alert\n→ 數位閱片清單", "#ede9fe", "#6d28d9")
    fc.box(7.7, 5.0, 3.6, 0.9, "AI + Follow-up 雙旗標\n→ 僅實體作業清單", "#dbeafe", "#1d4ed8")
    fc.box(2.3, 3.2, 3.2, 0.9, "綠/黃簽核 或\n改為人工鏡檢", "#f8fafc", "#64748b")
    fc.box(7.7, 3.2, 3.4, 0.9, "拉片完成時\n一併確認 AI", "#f8fafc", "#64748b")
    fc.box(5, 1.5, 4.5, 0.9, "[完成] AI 確認：膠囊變綠底打勾", "#dcfce7", "#15803d", 10)
    fc.arrow(5, 7.75, 5, 7.45)
    fc.arrow(4.2, 6.2, 2.3, 5.5, "僅 AI")
    fc.arrow(5.8, 6.2, 7.7, 5.5, "雙旗標")
    fc.arrow(2.3, 4.5, 2.3, 3.7)
    fc.arrow(7.7, 4.5, 7.7, 3.7)
    fc.arrow(2.3, 2.7, 4.0, 1.95)
    fc.arrow(7.7, 2.7, 6.0, 1.95)
    fc.save(path)


def chart_follow_up(path: Path):
    fc = FlowChart("圖 4　需拉片確認（Follow-up）生命週期", figsize=(10, 8))
    fc.box(5, 8.3, 4.0, 0.8, "[紅] 需拉片確認 待辦", "#fee2e2", "#b91c1c")
    fc.diamond(5, 6.9, 1.7, "如何產生？")
    fc.box(1.8, 5.2, 2.8, 0.9, "報告紅色橫幅\n改為人工鏡檢", "#fef3c7", "#b45309", 8)
    fc.box(5.0, 5.2, 2.6, 0.9, "閱片頁\nAdd Flag", "#f3e8ff", "#7c3aed", 8)
    fc.box(8.2, 5.2, 2.8, 0.9, "舊版 Manual Alert\n自動遷移", "#f1f5f9", "#64748b", 8)
    fc.box(5, 3.7, 4.2, 0.8, "實體作業待辦清單", "#dbeafe", "#1d4ed8")
    fc.diamond(5, 2.3, 1.6, "完成路徑？")
    fc.box(2.0, 0.9, 2.8, 0.8, "列表點膠囊", "#f8fafc", "#64748b", 8)
    fc.box(5.0, 0.9, 2.8, 0.8, "已拉片完成", "#f8fafc", "#64748b", 8)
    fc.box(8.0, 0.9, 2.8, 0.8, "開鎖強制簽核", "#f8fafc", "#64748b", 8)
    fc.note(5, 0.2, "完成後：entityStatusDone.Follow-up = true，膠囊變綠底打勾")
    fc.arrow(5, 7.85, 5, 7.55)
    fc.arrow(4.1, 6.2, 1.8, 5.7)
    fc.arrow(5, 6.2, 5, 5.7)
    fc.arrow(5.9, 6.2, 8.2, 5.7)
    fc.arrow(1.8, 4.7, 3.8, 4.1)
    fc.arrow(5, 4.7, 5, 4.15)
    fc.arrow(8.2, 4.7, 6.2, 4.1)
    fc.arrow(5, 3.25, 5, 3.0)
    fc.arrow(4.2, 1.8, 2.0, 1.35)
    fc.arrow(5, 1.8, 5, 1.35)
    fc.arrow(5.8, 1.8, 8.0, 1.35)
    fc.save(path)


def chart_report_banner(path: Path):
    fc = FlowChart("圖 5　報告核發 × 風險橫幅 × 膠囊轉換", figsize=(11, 9))
    fc.box(5.5, 8.5, 4.5, 0.8, "影像檢視 → 儲存並核發報告", "#f8fafc", "#64748b")
    fc.diamond(5.5, 7.0, 1.8, "風險橫幅？")
    fc.box(1.5, 5.3, 2.8, 0.9, "[綠] 未達留單", "#dcfce7", "#15803d", 8)
    fc.box(5.5, 5.3, 2.8, 0.9, "[黃] 延續性異常", "#fef9c3", "#a16207", 8)
    fc.box(9.5, 5.3, 2.8, 0.9, "[紅] 新發留單", "#fee2e2", "#b91c1c", 8)
    fc.box(1.5, 3.5, 2.8, 1.0, "確認並簽核\n[完成] DR + AI\n自動下一筆", "#ecfdf5", "#047857", 8)
    fc.box(5.5, 3.5, 2.8, 1.0, "確認並簽核\n[完成] DR + AI\n自動下一筆", "#ecfdf5", "#047857", 8)
    fc.box(9.5, 3.5, 2.8, 1.0, "改為人工鏡檢\n或開鎖強制簽核", "#fff7ed", "#9a3412", 8)
    fc.box(9.5, 1.5, 3.2, 1.0, "交接：[完成] DR + [紅] FU\n數位頁唯讀快照", "#ede9fe", "#6d28d9", 8)
    fc.arrow(5.5, 8.05, 5.5, 7.85)
    fc.arrow(4.7, 6.2, 1.5, 5.8, "綠")
    fc.arrow(5.5, 6.2, 5.5, 5.8, "黃")
    fc.arrow(6.3, 6.2, 9.5, 5.8, "紅")
    fc.arrow(1.5, 4.8, 1.5, 4.05)
    fc.arrow(5.5, 4.8, 5.5, 4.05)
    fc.arrow(9.5, 4.8, 9.5, 4.05)
    fc.arrow(9.5, 2.95, 9.5, 2.05, "改人工")
    fc.save(path)


def chart_list_routing(path: Path):
    fc = FlowChart("圖 6　清單歸屬：數位閱片 vs 實體作業", figsize=(10, 7))
    fc.box(2.5, 7.5, 3.8, 1.1, "數位閱片清單\n• DR 未完成\n• 僅 AI 待確認\n• Verified 篩選可見", "#ede9fe", "#6d28d9", 9)
    fc.box(7.5, 7.5, 3.8, 1.1, "實體作業清單\n• PLT / Follow-up 待辦\n• AI+Follow-up 雙旗標\n• 數位交接後拉片", "#dbeafe", "#1d4ed8", 9)
    fc.box(5, 5.0, 4.5, 0.9, "排除於數位清單", "#fef3c7", "#b45309", 10)
    fc.box(2.5, 2.8, 3.4, 0.9, "AI + Follow-up 雙旗標", "#fee2e2", "#b91c1c", 9)
    fc.box(7.5, 2.8, 3.6, 0.9, "數位已完成\n僅剩實體待辦", "#ffedd5", "#c2410c", 9)
    fc.arrow(5, 4.5, 2.5, 3.35)
    fc.arrow(5, 4.5, 7.5, 3.35)
    fc.arrow(2.5, 2.3, 2.5, 1.5)
    fc.arrow(7.5, 2.3, 7.5, 1.5)
    fc.note(5, 1.0, "雙模式可搭配狀態篩選、日期、單位、機台等條件微調")
    fc.save(path)


def chart_readonly(path: Path):
    fc = FlowChart("圖 7　進入閱片 × 唯讀規則", figsize=(10, 8))
    fc.box(5, 8.3, 3.6, 0.8, "列表：進入閱片 / 唯讀檢視", "#f8fafc", "#64748b")
    fc.diamond(5, 6.9, 1.6, "locked？")
    fc.box(2.0, 5.2, 3.0, 0.9, "[鎖定] 唯讀\n他人編輯中", "#e5e7eb", "#374151", 9)
    fc.diamond(7.8, 6.9, 1.6, "唯讀？")
    fc.box(7.8, 5.0, 3.2, 0.8, "進入閱片（可編輯）", "#dcfce7", "#15803d", 9)
    fc.diamond(5, 4.0, 1.6, "唯讀原因？")
    fc.box(2.0, 2.3, 3.2, 0.9, "已簽核結案\nsignedOff=true", "#fff7ed", "#9a3412", 8)
    fc.box(5.0, 2.3, 3.2, 0.9, "交接 Follow-up\n數位快照唯讀", "#fff7ed", "#9a3412", 8)
    fc.box(8.0, 2.3, 3.0, 0.8, "右鍵退回 DR\n可再編輯", "#e0f2fe", "#0369a1", 8)
    fc.note(5, 0.8, "唯讀時：可瀏覽縮放，不可改分類、Add Flag、核發報告")
    fc.arrow(5, 7.85, 5, 7.65)
    fc.arrow(4.2, 6.4, 2.0, 5.7, "是")
    fc.arrow(5.8, 6.4, 7.8, 7.2, "否")
    fc.arrow(7.8, 6.2, 7.8, 5.45, "否")
    fc.arrow(7.8, 6.2, 5, 4.75, "是")
    fc.arrow(4.2, 3.5, 2.0, 2.8)
    fc.arrow(5, 3.5, 5, 2.8)
    fc.arrow(5.8, 3.5, 8.0, 2.8)
    fc.save(path)


def generate_charts() -> dict[str, Path]:
    setup_matplotlib_font()
    ASSETS.mkdir(parents=True, exist_ok=True)
    charts = {
        "overview": ASSETS / "01_overview.png",
        "digital_review": ASSETS / "02_digital_review.png",
        "ai_alert": ASSETS / "03_ai_alert.png",
        "follow_up": ASSETS / "04_follow_up.png",
        "report_banner": ASSETS / "05_report_banner.png",
        "list_routing": ASSETS / "06_list_routing.png",
        "readonly": ASSETS / "07_readonly.png",
    }
    chart_overview(charts["overview"])
    chart_digital_review(charts["digital_review"])
    chart_ai_alert(charts["ai_alert"])
    chart_follow_up(charts["follow_up"])
    chart_report_banner(charts["report_banner"])
    chart_list_routing(charts["list_routing"])
    chart_readonly(charts["readonly"])
    return charts


def build_document(charts: dict[str, Path]):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("血球分類軟體介面專案")
    set_run_font(run, 16, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run("狀態膠囊流程說明"), 20, bold=True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta.add_run("含流程圖、文字說明與速查表"), 11, color=(100, 100, 100))
    doc.add_page_break()

    add_title(doc, "一、文件目的", 1)
    add_para(
        doc,
        "本文件整理血球分類軟體中「狀態膠囊」的各種情境，包含雙流程分流（數位閱片／實體作業）、"
        "各膠囊的待辦與完成條件、報告核發時的風險橫幅轉換、唯讀規則，以及整體流程完成判定。"
        "內容依據 common.js、檢體管理.html、image-review.js、report-issue.js 之現行邏輯整理。"
    )

    add_title(doc, "二、系統總覽", 1)
    add_para(doc, "系統以兩條平行流程管理檢體：數位流程由鏡檢醫檢師負責，實體流程由核發醫檢師負責。兩邊都通過後，檢體視為整體完成。")
    add_image(doc, charts["overview"], "圖 1　系統總覽：雙流程與整體結案")
    add_table(
        doc,
        ["流程邊", "無須處理條件", "須完成條件"],
        [
            ["數位", "狀態不含 Digital Review", "workflowDone.digitalReview = true"],
            ["數位", "狀態不含 AI Alert", "workflowDone.aiAlertConfirmed = true"],
            ["實體", "狀態不含 PLT Check / Follow-up", "各膠囊 entityStatusDone[key] = true"],
        ],
    )

    add_title(doc, "三、膠囊外觀對照", 1)
    add_para(doc, "下列對照圖展示各膠囊「待辦」與「完成」時的視覺樣式（色碼與 common.js STATUS_STYLES 一致）。")
    add_image(doc, CAPSULE_STYLE_IMG, "圖 8　狀態膠囊樣式對照圖", width_in=6.5)
    add_table(
        doc,
        ["膠囊", "待辦外觀", "完成外觀", "誰能點擊切換", "完成觸發"],
        [
            ["數位閱片", "紫底", "綠底 + ✓", "不可點擊", "簽核 / 改為人工鏡檢"],
            ["AI 分類警示", "橘底", "綠底 + ✓", "不可點擊", "簽核 / 改為人工鏡檢 / 雙旗標拉片完成"],
            ["需拉片確認", "紅底", "綠底 + ✓", "列表點擊（實體模式）", "點膠囊 / 已拉片完成 / 開鎖簽核"],
            ["血小板確認", "藍底", "綠底 + ✓", "列表點擊（實體模式）", "點膠囊"],
            ["鎖定中", "灰底 + 🔒", "—", "不可操作", "資料庫 locked 欄位"],
            ["Verified", "不顯示膠囊", "—", "篩選勾選", "整體流程完成"],
        ],
    )

    add_title(doc, "四、各膠囊生命週期", 1)

    add_title(doc, "4.1 數位閱片（Digital Review）", 2)
    add_bullets(
        doc,
        [
            "待辦時出現在數位閱片清單，可進入影像檢視頁編輯細胞分類。",
            "綠／黃橫幅簽核：標記完成並簽核結案（digitalReviewSignedOff = true），之後唯讀。",
            "紅色橫幅「改為人工鏡檢」：數位閱片標完成但未簽核結案，並加入需拉片確認；數位頁進入交接唯讀。",
            "右鍵「退回 Digital Review」可重開數位流程，清除簽核與編輯快照。",
        ],
    )
    add_image(doc, charts["digital_review"], "圖 2　數位閱片生命週期")

    add_title(doc, "4.2 AI 分類警示（AI Alert）", 2)
    add_bullets(
        doc,
        [
            "僅 AI Alert：出現在數位閱片清單，待綠／黃簽核或改為人工鏡檢時確認。",
            "AI + Follow-up 雙旗標：僅出現在實體作業清單；拉片完成時一併確認 AI。",
            "完成後膠囊轉為綠底打勾（aiAlertConfirmed = true）。",
        ],
    )
    add_image(doc, charts["ai_alert"], "圖 3　AI 分類警示生命週期")

    add_title(doc, "4.3 需拉片確認（Follow-up）", 2)
    add_bullets(
        doc,
        [
            "舊版 Manual Alert 會自動併入 Follow-up。",
            "常見來源：報告紅色橫幅「改為人工鏡檢」、閱片頁 Add Flag。",
            "於實體作業模式處理：列表點膠囊、報告頁「已拉片完成」，或開鎖強制簽核。",
        ],
    )
    add_image(doc, charts["follow_up"], "圖 4　需拉片確認生命週期")

    add_title(doc, "4.4 血小板確認（PLT Check）", 2)
    add_bullets(
        doc,
        [
            "屬實體作業膠囊，出現在實體作業清單。",
            "於列表直接點擊膠囊可切換完成狀態。",
            "完成後 entityStatusDone['PLT Check'] = true。",
        ],
    )

    add_title(doc, "4.5 鎖定中（Locked）與已完成（Verified）", 2)
    add_bullets(
        doc,
        [
            "Locked：他人編輯中，列表斜紋背景，無法進入編輯閱片，僅能唯讀檢視。",
            "Verified：不作為膠囊顯示，僅作篩選條件；整體流程完成後時效歸零、列表沉底。",
        ],
    )

    add_title(doc, "五、報告核發與風險橫幅", 1)
    add_para(
        doc,
        "報告核發頁依留單門檻與前次報告比對，顯示綠／黃／紅三色橫幅。若醫檢師曾修改細胞分類，"
        "風險判讀以人員編輯數值為準。紅色新發留單時，簽核預設鎖定，可改為人工鏡檢或開鎖強制簽核。"
    )
    add_image(doc, charts["report_banner"], "圖 5　報告核發 × 風險橫幅 × 膠囊轉換")
    add_table(
        doc,
        ["橫幅", "意義", "簽核行為", "膠囊影響"],
        [
            ["綠色", "未達留單標準", "可直接簽核，可自動下一筆", "✅ DR；✅ AI（若有）"],
            ["黃色", "延續性異常", "可直接簽核，可自動下一筆", "✅ DR；✅ AI（若有）"],
            ["紅色", "新發留單", "預設鎖定；可改人工或開鎖", "改人工：✅ DR + 🔴 FU；開鎖：強制結案"],
        ],
    )

    add_title(doc, "六、清單歸屬規則", 1)
    add_image(doc, charts["list_routing"], "圖 6　清單歸屬：數位閱片 vs 實體作業")
    add_bullets(
        doc,
        [
            "數位閱片模式：列出 Digital Review 未完成，或僅 AI Alert 待確認者。",
            "排除：AI + Follow-up 雙旗標；數位已完成但僅剩實體待辦（如 PLT Check）。",
            "實體作業模式：列出 PLT Check、Follow-up、AI+Follow-up 等待辦者。",
            "整體已完成者，勾選 Verified 篩選時仍可見。",
        ],
    )

    add_title(doc, "七、進入閱片與唯讀規則", 1)
    add_image(doc, charts["readonly"], "圖 7　進入閱片 × 唯讀規則")
    add_table(
        doc,
        ["情境", "列表按鈕", "閱片頁", "能否重開編輯"],
        [
            ["數位閱片進行中", "進入閱片", "可編輯", "—"],
            ["改為人工鏡檢後", "唯讀檢視", "交接快照唯讀", "右鍵退回 DR"],
            ["數位已簽核結案", "唯讀檢視", "簽核結案唯讀", "右鍵退回 DR"],
            ["Locked", "進入閱片停用", "唯讀", "等解鎖"],
        ],
    )

    add_title(doc, "八、常見組合情境", 1)
    add_table(
        doc,
        ["情境", "膠囊狀態", "所在清單", "閱片模式", "下一步"],
        [
            ["純數位閱片", "🟣 DR", "數位", "可編輯", "簽核 → ✅ DR"],
            ["僅 AI 警示", "🟠 AI", "數位", "可進報告", "簽核 → ✅ AI"],
            ["純血小板", "🔵 PLT", "實體", "不進閱片", "點膠囊 → ✅"],
            ["紅色留單改人工", "✅ DR + 🔴 FU", "實體", "唯讀快照", "拉片 → 簽核"],
            ["AI + 拉片雙旗標", "🟠 AI + 🔴 FU", "實體", "—", "拉片完成同時 ✅ AI"],
            ["全部完成", "全綠 ✓", "Verified 篩選", "唯讀", "時效 = 0"],
            ["他人鎖定", "🔒", "任一", "唯讀", "等解鎖"],
        ],
    )

    add_title(doc, "九、技術備註", 1)
    add_bullets(
        doc,
        [
            "workflowDone 欄位：digitalReview、digitalReviewSignedOff、aiAlertConfirmed、entityStatusDone。",
            "狀態覆寫儲存於 localStorage 鍵 blood-morphology-specimen-status。",
            "重新登入 Demo 會清除狀態覆寫，還原 database.js 初始資料。",
            "本文件流程圖為示意簡化版，實際判定以程式碼為準。",
        ],
    )

    doc.save(str(OUTPUT))
    return OUTPUT


def main():
    print("產生流程圖…")
    charts = generate_charts()
    print("組裝 Word 文件…")
    out = build_document(charts)
    print(f"完成：{out}")


if __name__ == "__main__":
    main()
